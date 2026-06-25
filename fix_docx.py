import docx

doc = docx.Document('public/FORMATO DE PERMISOS - REPOSOS.docx')

# 1. First, let's find the small 3-cell table at the top right to put {d_gen}, {m_gen}, {a_gen}.
for t in doc.tables:
    # A small table with 1 row and 3 columns? Or 1 row and more cols where only last 3 are used?
    if len(t.rows) == 1 and len(t.rows[0].cells) == 3:
        t.rows[0].cells[0].text = '{d_gen}'
        t.rows[0].cells[1].text = '{m_gen}'
        t.rows[0].cells[2].text = '{a_gen}'

# Fix Table 0
t0 = doc.tables[0]
# Cédula
for cell in t0.rows[3].cells[5:8]: cell.text = ''
t0.rows[3].cells[5].text = '{cedula}'
# Código
t0.rows[3].cells[13].text = '{codigo}'
# Cargo
for cell in t0.rows[3].cells[8:13]: cell.text = ''
t0.rows[3].cells[8].text = '{cargo}'

# Clear out the erroneous {cargo} in the signature lines
for cell in t0.rows[6].cells:
    if '{cargo}' in cell.text:
        cell.text = cell.text.replace('{cargo}', '')

# Fix Observación (Row 7)
for cell in t0.rows[7].cells:
    cell.text = ''
t0.rows[7].cells[1].text = '6. OBSERVACION\nSE ANEXA CONSTANCIA MÉDICA DE REPOSO POR {anexo}\nObservación: {observacion}'

# Fix Table 1 (Fechas desde/hasta en formato DIA MES AÑO)
t1 = doc.tables[1]
# Clear out any existing {fecha_inicio} etc.
for r in t1.rows:
    for c in r.cells:
        if '{fecha_inicio}' in c.text: c.text = c.text.replace('{fecha_inicio}', '')
        if '{fecha_culminacion}' in c.text: c.text = c.text.replace('{fecha_culminacion}', '')

# Put the split dates in Row 4 (since Row 3 in my previous dump had the headers DIA MES AÑO in the user's pic, wait:
# The user pic shows:
# Row 0: 7. Duración del Permiso
# Row 1: 7.1 HORA(S) | 7.2 DIA(S)
# Row 2: DESDE | HASTA | DESDE | HASTA
# Row 3: (blank) | (blank) | DIA | MES | AÑO | DIA | MES | AÑO
# Row 4: values...

# Let's just hardcode the cells for Row 4 based on the columns:
# It seems column index 2 is DIA desde, 3 is MES desde, 4 is AÑO desde.
# Column index 5 is DIA hasta, 6 is MES hasta, 7 is AÑO hasta.
# Let's clear row 4 cells 2 to 7
for cell in t1.rows[4].cells[2:8]: cell.text = ''
t1.rows[4].cells[2].text = '{d_desde}'
t1.rows[4].cells[3].text = '{m_desde}'
t1.rows[4].cells[4].text = '{a_desde}'
t1.rows[4].cells[5].text = '{d_hasta}'
t1.rows[4].cells[6].text = '{m_hasta}'
t1.rows[4].cells[7].text = '{a_hasta}'

# Let's also check if Row 3 needs anything removed
for cell in t1.rows[3].cells[2:8]:
    if '{fecha' in cell.text:
        cell.text = ''

# Save
doc.save('public/FORMATO DE PERMISOS - REPOSOS.docx')
print("Modificado con éxito")
