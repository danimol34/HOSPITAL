import docx

doc = docx.Document('public/FORMATO DE PERMISOS - REPOSOS.docx')

def scan_tables(tables, source):
    for i, t in enumerate(tables):
        print(f"[{source}] Table {i}: {len(t.rows)} rows")
        for r in t.rows:
            row_texts = [c.text.strip() for c in r.cells]
            print(f"  Row cols={len(r.cells)}: {row_texts}")
            if len(r.cells) == 3 or (len(r.cells) >= 3 and all(c=='' for c in row_texts[-3:])):
                print("  => Potential match for date boxes!")

print("Scanning main body...")
scan_tables(doc.tables, "Body")

for section in doc.sections:
    print("Scanning header...")
    scan_tables(section.header.tables, "Header")
    print("Scanning footer...")
    scan_tables(section.footer.tables, "Footer")
