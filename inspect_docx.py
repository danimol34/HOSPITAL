import docx

def inspect_docx(filename):
    doc = docx.Document(filename)
    print(f"--- Document: {filename} ---")
    print("PARAGRAPHS:")
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            print(f"P{i}: {p.text}")
    print("\nTABLES:")
    for i, table in enumerate(doc.tables):
        print(f"\nTable {i}:")
        for r_idx, row in enumerate(table.rows):
            row_texts = []
            for c_idx, cell in enumerate(row.cells):
                text = cell.text.replace('\n', ' ').strip()
                row_texts.append(f"[{c_idx}] {text}")
            print(f"  Row {r_idx}: " + " | ".join(row_texts))

inspect_docx('public/FORMATO DE PERMISOS - REPOSOS.docx')
